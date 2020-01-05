import json
from docx import Document

def demo():
    document = Document()

    document.add_heading("Document Title", 0)

    p = document.add_paragraph("A plain paragraph having some ")
    p.add_run("bold").bold = True
    p.add_run(" and some ")
    p.add_run("italic.").italic = True

    document.save("./documents/demo.docx")

def parseReport(data, date):
    document = Document()
    document.add_heading("Sport Scan Daily Brief", 0)
    document.add_paragraph(date).italic = True


    # for team in data['teams']:
    #     print(team['name'])
    #     for title in team['titles']:
    #         print(title)

    for team in data["teams"]:
        document.add_heading(team['name'], 1)
        for title in team["titles"]:
            document.add_paragraph(title)

    document.save("./documents/report.docx")
    return document

# data = '{ "teams": [{  "name": "NY Yankees",  "titles": ["article1", "article2", "article3"] }, { "name": "Boston Red Sox",  "titles": ["article1", "article2", "article3"]  }  ]}'

# parseReport(json.loads(data), "12/26/1996")